from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt
from datetime import datetime
import pandas as pd
import io

# Define the file paths
template_path = 'Template.docx'
csv_path = 'data.csv'

# Load the template Word document
def load_template(template_path):
    try:
        return DocxTemplate(template_path)
    except Exception as e:
        print(f"Error loading template: {e}")
        exit(1)

# Create a dictionary with the content to be rendered in the template
def create_content():
    return {
        'name': 'Dan Bhatia',
        's_time': '1045 hrs. AEST',
        'e_time': '1400 hrs. AEST',
        'i_name': 'Robin and John',
        'date': datetime.today().strftime('%Y%m%d_01')
    }

# Render the template and save to an in-memory file
def render_template(doc, content):
    try:
        doc.render(content)
        fake_file = io.BytesIO()
        doc.save(fake_file)
        fake_file.seek(0)
        return fake_file
    except Exception as e:
        print(f"Error rendering template: {e}")
        exit(1)

# Read data from CSV file
def read_csv(csv_path):
    try:
        return pd.read_csv(csv_path)
    except Exception as e:
        print(f"Error reading CSV file: {e}")
        exit(1)

# Function to add a bullet point with custom indentation
def add_bullet_point(paragraph, text, level=0):
    run = paragraph.add_run(f"• {text}")
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    paragraph.paragraph_format.left_indent = Pt(14.4 * level)

# Add job title and details from CSV to the document
def add_job_details(document, df):
    for i, row in df.iterrows():
        # Add job title as bold heading
        job_title = document.add_paragraph(f"{i + 1}. {row['Job_Title']}:")
        job_title.runs[0].bold = True
        
        # Add name, surname, age, and date with indentation
        for field in ['Name', 'Surname', 'Age', 'Date']:
            para = document.add_paragraph()
            add_bullet_point(para, f"{field} : {row[field]}", 1)
        
        if i < len(df) - 1:  # To avoid adding a page break after the last entry
            document.add_page_break()

# Function to add a new page at the end with the title - "All Files Screenshots:"
def add_screenshots_page(document):
    document.add_page_break()  # Add a page break to ensure it's on a new page
    paragraph = document.add_paragraph()
    run = paragraph.add_run("All Files Screenshots:")
    run.bold = True
    run.font.size = Pt(14)

# Save the final document with the added paragraphs
def save_document(document, date):
    output_path = f"{date}.docx"
    try:
        document.save(output_path)
        print(f"Document saved as {output_path}")
    except Exception as e:
        print(f"Error saving document: {e}")
        exit(1)

def main():
    doc = load_template(template_path)
    content = create_content()
    fake_file = render_template(doc, content)
    document = Document(fake_file)
    df = read_csv(csv_path)
    add_job_details(document, df)
    add_screenshots_page(document)
    save_document(document, content['date'])

if __name__ == "__main__":
    main()
